"""
报告导出与载荷数据导出模块
功能1: 载荷数据导出（多通道 → Excel/CSV/TXT）
功能2: 分析结果报告导出（PDF/Word）
"""
import os
import io
from datetime import datetime

import numpy as np
import pandas as pd

# matplotlib 中文支持
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import font_manager

# 中文字体配置
for _f in ["Microsoft YaHei", "SimHei", "SimSun", "PingFang SC", "Noto Sans CJK SC"]:
    try:
        font_manager.findfont(_f, fallback_to_default=False)
        plt.rcParams["font.sans-serif"] = [_f]
        break
    except Exception:
        continue
plt.rcParams["axes.unicode_minus"] = False


# ============================================================
# 功能1: 载荷数据导出
# ============================================================
def export_channels(df, channels, fmt="excel", output_dir=None, prefix="载荷数据", filename=None):
    """
    导出指定通道随时间变化的数据表
    df: 已加载的DataFrame（含time列）
    channels: 通道列名列表
    fmt: excel / csv / txt
    output_dir: 保存目录（可选）
    prefix: 文件名前缀（可选）
    filename: 自定义文件名（含扩展名，可选，优先于prefix）
    返回: (文件路径, 文件名)
    """
    if df is None or len(channels) == 0:
        raise ValueError("无数据或未选择通道")

    # 检测时间列（优先用配置，兼容空格/下划线变体）
    time_col = _find_time_col(df)
    cols = [c for c in channels if c in df.columns]
    if not cols:
        raise ValueError("所选通道不在数据中")

    out_df = pd.DataFrame()
    if time_col and time_col in df.columns:
        out_df["time [s]"] = df[time_col].values
    for c in cols:
        out_df[c] = df[c].values

    output_dir = output_dir or os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)

    fmt = fmt.lower()
    ext_map = {"xlsx": "xlsx", "excel": "xlsx", "xls": "xlsx",
               "csv": "csv", "txt": "txt", "dat": "txt"}
    ext = ext_map.get(fmt, "xlsx")

    if filename:
        # 确保扩展名正确
        if not filename.lower().endswith("." + ext):
            filename = filename + "." + ext
        fname = filename
    else:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        fname = f"{prefix}_{ts}.{ext}"

    filepath = os.path.join(output_dir, fname)

    if fmt in ("xlsx", "excel", "xls"):
        out_df.to_excel(filepath, index=False)
    elif fmt == "csv":
        out_df.to_csv(filepath, index=False, encoding="utf-8-sig")
    elif fmt in ("txt", "dat"):
        out_df.to_csv(filepath, index=False, sep="\t", encoding="utf-8")
    else:
        raise ValueError(f"不支持的格式: {fmt}")

    return filepath, fname


# ============================================================
# 工具: 渲染时序图（matplotlib → PNG bytes）
# ============================================================
def render_channel_chart(df, channel, max_points=500, figsize=(8, 3.2)):
    """渲染单通道时序图，返回PNG bytes"""
    time_col = config_time_col(df)
    t = df[time_col].values
    vals = df[channel].values
    step = max(1, len(t) // max_points)
    t_s = t[::step]
    v_s = vals[::step]

    fig, ax = plt.subplots(figsize=figsize)
    ax.plot(t_s, v_s, lw=0.8, color="#1565c0")
    mean_val = np.mean(vals)
    ax.axhline(mean_val, color="#ef6c00", ls="--", lw=1, label=f"均值={mean_val:.2f}")
    ax.set_title(channel, fontsize=10)
    ax.set_xlabel("Time [s]", fontsize=9)
    ax.set_ylabel("Value", fontsize=9)
    ax.tick_params(labelsize=8)
    ax.legend(fontsize=8)
    ax.grid(alpha=0.3)
    fig.tight_layout()
    buf = io.BytesIO()
    fig.savefig(buf, format="png", dpi=110)
    plt.close(fig)
    buf.seek(0)
    return buf.getvalue()


def _find_time_col(df):
    """检测DataFrame中的时间列（兼容空格/下划线变体）"""
    # 优先用配置
    try:
        from config import config
        tc = config.global_ch.time_col
        if tc in df.columns:
            return tc
    except Exception:
        pass
    for c in ["time", "Time [s]", "Time_[s]", "Time", "Time [s] ", "time [s]", "time_[s]"]:
        if c in df.columns:
            return c
    return None


def config_time_col(df):
    tc = _find_time_col(df)
    return tc if tc else df.columns[0]


# 常用通道定义（B2修复：pattern 均为正则，用 re.search 匹配，兼容列名变体）
COMMON_CHANNEL_PATTERNS = [
    ("风轮推力", [r"Momentary_Aerodynamic_Thrust", r"Momentary_Aerodynamic thrust"]),
    ("风轮扭矩", [r"Momentary_Aerodynamic_Torque", r"Momentary_Aerodynamic torque"]),
    ("入流风速", [r"Abs_Meas\._Wind_Vel\._at_Hub", r"Abs_Meas\._Wind_Velocity_at_Hub",
                 r"Abs.*Wind.*Hub", r"Abs_Inflow_Vel.*Hub", r"Inflow.*Vel.*Hub"]),
    ("气动功率", [r"Momentary_Aerodynamic_Power"]),
    ("转速", [r"Rotational_Speed", r"Rotational speed"]),
]


def find_common_channels(df):
    """
    在df中查找常用通道（正则匹配，与前端 COMMON_CHANNELS 规则保持一致）
    返回在原始列顺序中最先命中的列名
    """
    import re
    found = []
    for label, patterns in COMMON_CHANNEL_PATTERNS:
        for col in df.columns:
            if any(re.search(p, col, re.IGNORECASE) for p in patterns):
                found.append(col)
                break
    return found


# ============================================================
# 功能2: 分析结果报告导出
# ============================================================
def generate_report(results, df, fmt="pdf", output_dir=None, calc_time=None, filename=None):
    """
    生成分析结果报告（PDF或Word）
    results: STATE["results"]（序列化后的完整结果）
    df: 已加载的DataFrame
    fmt: pdf / word
    output_dir: 保存目录（可选）
    calc_time: 计算耗时
    filename: 自定义文件名（含扩展名，可选）
    返回: (文件路径, 文件名)
    """
    output_dir = output_dir or os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = "垂直轴风轮载荷分析结果"

    # 构建统一内容模型
    content = _build_report_content(results, df, calc_time)

    if fmt.lower() in ("pdf",):
        ext = "pdf"
        if filename:
            fname = filename if filename.lower().endswith(".pdf") else filename + ".pdf"
        else:
            fname = f"{base_name}_{ts}.pdf"
        filepath = os.path.join(output_dir, fname)
        _write_pdf(filepath, content)
    elif fmt.lower() in ("word", "docx", "doc"):
        ext = "docx"
        if filename:
            fname = filename if filename.lower().endswith(".docx") else filename + ".docx"
        else:
            fname = f"{base_name}_{ts}.docx"
        filepath = os.path.join(output_dir, fname)
        _write_word(filepath, content)
    else:
        raise ValueError(f"不支持的格式: {fmt}")

    return filepath, fname


def _build_report_content(results, df, calc_time=None):
    """构建报告内容模型，供PDF/Word共用"""
    content = {
        "title": "垂直轴风轮载荷分析结果",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "calc_time": calc_time or "--",
        "sections": [],
    }

    fi = (results or {}).get("file_info", {})
    info_rows = [
        ("文件名", fi.get("filename", "--")),
        ("文件大小", f"{fi.get('size_mb', 0):.2f} MB" if isinstance(fi.get("size_mb"), (int, float)) else "--"),
        ("工况类型", fi.get("condition_type", "--")),
        ("工况名称", fi.get("condition_name", "--")),
        ("创建时间", fi.get("creation_time", "--")),
        ("数据", f"{fi.get('num_rows', 0)} 行 x {fi.get('num_cols', 0)} 列"),
        ("叶片数", fi.get("num_blades", "--")),
        ("面板数", fi.get("num_panels", "--")),
        ("采样率", f"{fi.get('sample_rate', 0)} Hz"),
        ("时长", f"{fi.get('duration', 0)} s"),
    ]
    # 通道映射信息
    mapper = (results or {}).get("mapper_info", {})
    if mapper:
        info_rows.append(("通道映射", f"{mapper.get('num_global', 0)}全局列, {mapper.get('num_blades', 0)}叶片, "
                                       f"{mapper.get('num_total_loads', 0)}总载荷列, {mapper.get('num_panel', 0)}面板级列"))
    content["sections"].append({
        "title": "输入文件信息",
        "type": "info_table",
        "rows": info_rows,
    })

    # ===== 第一章: 结论（风轮极限载荷极值汇总）=====
    ext = (results or {}).get("extreme") or (results or {}).get("eog") or {}
    wheel = ext.get("wheel", {})
    conclusion_rows = []
    if wheel.get("thrust_stats"):
        ts_ = wheel["thrust_stats"]
        conclusion_rows.append(("风轮推力 (Momentary Aerodynamic Thrust [N])",
                                f"最大={ts_.get('max', '--')}", f"最小={ts_.get('min', '--')}",
                                f"均值={ts_.get('mean', '--')}",
                                f"最大时刻={wheel.get('thrust_max_time', '--')} s"))
    if wheel.get("torque_stats"):
        ts_ = wheel["torque_stats"]
        conclusion_rows.append(("风轮扭矩 (Momentary Aerodynamic Torque [Nm])",
                                f"最大={ts_.get('max', '--')}", f"最小={ts_.get('min', '--')}",
                                f"均值={ts_.get('mean', '--')}",
                                f"最大时刻={wheel.get('torque_max_time', '--')} s"))
    if wheel.get("power_stats"):
        ts_ = wheel["power_stats"]
        conclusion_rows.append(("风轮功率 (Momentary Aerodynamic Power [W])",
                                f"最大={ts_.get('max', '--')}", f"最小={ts_.get('min', '--')}",
                                f"均值={ts_.get('mean', '--')}",
                                f"最大时刻={wheel.get('power_max_time', '--')} s"))
    if wheel.get("overturning_moment_stats"):
        ts_ = wheel["overturning_moment_stats"]
        conclusion_rows.append(("倾覆弯矩 [Nm]",
                                f"最大={ts_.get('max', '--')}", f"最小={ts_.get('min', '--')}",
                                f"均值={ts_.get('mean', '--')}", "--"))
    if wheel.get("critical_blade"):
        conclusion_rows.append(("最不利叶片", wheel.get("critical_blade", "--"),
                                f"峰值={wheel.get('critical_blade_load', '--')} N", "--", "--"))
    if wheel.get("combined_peak_imbalance") is not None:
        conclusion_rows.append(("合成峰值不平衡度", f"{wheel.get('combined_peak_imbalance', '--')}", "--", "--", "--"))

    content["sections"].insert(0, {
        "title": "一、结论：风轮极限载荷极值汇总",
        "type": "table",
        "headers": ["载荷项", "最大值", "最小值", "均值", "最大发生时刻"],
        "rows": conclusion_rows,
    })

    # ===== 常用通道时序视图 =====
    common_chs = find_common_channels(df) if df is not None else []
    if common_chs:
        charts = []
        for ch in common_chs[:5]:
            try:
                charts.append({"title": ch, "image": render_channel_chart(df, ch)})
            except Exception:
                pass
        content["sections"].append({
            "title": "二、常用通道时序视图",
            "type": "charts",
            "charts": charts,
        })

    # ===== 极限载荷分析 =====
    if ext.get("summary"):
        content["sections"].append({
            "title": "三、叶片极限载荷分析",
            "type": "table",
            "headers": list(ext["summary"][0].keys()) if ext["summary"] else [],
            "rows": [list(r.values()) for r in ext["summary"]],
        })

    # ===== 疲劳分析 =====
    fat = (results or {}).get("fatigue", {})
    if fat.get("summary"):
        content["sections"].append({
            "title": "四、疲劳损伤分析",
            "type": "table",
            "headers": list(fat["summary"][0].keys()) if fat["summary"] else [],
            "rows": [list(r.values()) for r in fat["summary"]],
        })
    if fat.get("wheel"):
        fat_rows = [(k, str(v)) for k, v in fat["wheel"].items()]
        content["sections"].append({
            "title": "四、疲劳损伤分析（整机汇总）",
            "type": "info_table",
            "rows": fat_rows,
        })

    # ===== 叶片合成 =====
    syn = (results or {}).get("synthesis", {})
    if syn:
        syn_rows = []
        if syn.get("imbalance_normal"):
            s = syn["imbalance_normal"]
            syn_rows.append(("法向载荷不平衡度-均值", str(s.get("imbalance_mean", "--"))))
            syn_rows.append(("法向载荷不平衡度-最大", str(s.get("imbalance_max", "--"))))
        if syn.get("imbalance_tangential"):
            s = syn["imbalance_tangential"]
            syn_rows.append(("切向载荷不平衡度-均值", str(s.get("imbalance_mean", "--"))))
            syn_rows.append(("切向载荷不平衡度-最大", str(s.get("imbalance_max", "--"))))
        if syn.get("pulsation"):
            p = syn["pulsation"]
            syn_rows.append(("旋转脉动-平均转速", f"{p.get('avg_rpm', '--')} rpm"))
            syn_rows.append(("旋转脉动-脉动率", str(p.get("pulsation_ratio", "--"))))
        if syn_rows:
            content["sections"].append({
                "title": "五、多叶片载荷合成与不平衡度",
                "type": "info_table",
                "rows": syn_rows,
            })

    # ===== 塔顶载荷分析 =====
    tt = (results or {}).get("tower_top", {})
    if tt and tt.get("extreme"):
        tt_rows = []
        for name, ex in tt["extreme"].items():
            fa = tt.get("fatigue", {}).get(name, {})
            tt_rows.append((
                name,
                f"最大={ex.get('max', '--')}", f"最小={ex.get('min', '--')}",
                f"均值={ex.get('mean', '--')}",
                f"最大时刻={ex.get('max_time', '--')} s",
                f"等效疲劳={fa.get('del', '--')} (m={fa.get('m', '--')})",
            ))
        content["sections"].append({
            "title": "六、塔顶载荷分析",
            "type": "table",
            "headers": ["通道", "最大值", "最小值", "均值", "最大发生时刻", "等效疲劳载荷"],
            "rows": tt_rows,
        })

    # ===== 叶片校核载荷（垂直轴：侧面固定，法向+切向+合成）=====
    bc = (results or {}).get("blade_check", {})
    if bc:
        for bid, benv in bc.items():
            rows = []
            for r in benv.get("extreme", []):
                rows.append((r.get("通道", ""),
                             f"上限={r.get('上限', '--')} ({r.get('上限来源', '--')}@{r.get('上限时刻', '--')}s)",
                             f"下限={r.get('下限', '--')} ({r.get('下限来源', '--')}@{r.get('下限时刻', '--')}s)"))
            for r in benv.get("fatigue", []):
                rows.append((r.get("通道", "") + " 等效疲劳(m=%s)" % r.get("Wöhler指数m", "--"),
                             f"{r.get('等效疲劳载荷', '--')}", "--"))
            if rows:
                content["sections"].append({
                    "title": f"七、叶片校核载荷 - {bid}",
                    "type": "table",
                    "headers": ["校核项", "极限/疲劳值", "备注"],
                    "rows": rows,
                })

    # ===== 计算信息 =====
    content["sections"].append({
        "title": "八、计算信息",
        "type": "info_table",
        "rows": [
            ("报告生成时间", content["generated_at"]),
            ("计算耗时", str(content["calc_time"])),
        ],
    })

    return content


# ============================================================
# PDF生成（reportlab）
# ============================================================
def _write_pdf(filepath, content):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.units import mm
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import (SimpleDocTemplate, Paragraph, Spacer,
                                    Table, TableStyle, Image, PageBreak)
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont

    # 注册中文字体（reportlab内置CID字体，无需额外字体文件）
    CJK_FONT = "STSong-Light"
    try:
        pdfmetrics.registerFont(UnicodeCIDFont(CJK_FONT))
    except Exception:
        CJK_FONT = "Helvetica"

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle("TitleCN", parent=styles["Title"], fontSize=20,
                                 leading=26, alignment=TA_CENTER, fontName=CJK_FONT)
    h1_style = ParagraphStyle("H1CN", parent=styles["Heading1"], fontSize=15,
                              leading=20, spaceBefore=14, spaceAfter=8, fontName=CJK_FONT)
    body_style = ParagraphStyle("BodyCN", parent=styles["BodyText"], fontSize=10,
                                leading=14, fontName=CJK_FONT)
    cell_style = ParagraphStyle("CellCN", fontName=CJK_FONT, fontSize=8, leading=10,
                                alignment=TA_LEFT)

    # 页面可用宽度（A4 210mm - 左右边距 18mm*2 = 174mm）
    AVAIL_WIDTH = 174 * mm

    doc = SimpleDocTemplate(filepath, pagesize=A4,
                            leftMargin=18*mm, rightMargin=18*mm,
                            topMargin=16*mm, bottomMargin=16*mm,
                            title=content["title"])
    story = []

    # 标题
    story.append(Paragraph(content["title"], title_style))
    story.append(Spacer(1, 6))
    story.append(Paragraph(f"报告生成时间: {content['generated_at']}", body_style))
    story.append(Spacer(1, 12))

    def _fmt_cell(v):
        """格式化单元格内容"""
        if isinstance(v, (int, float)):
            if abs(v) < 0.01 and v != 0:
                return f"{v:.3e}"
            return f"{v:.4g}"
        return str(v)

    for sec in content["sections"]:
        story.append(Paragraph(sec["title"], h1_style))
        stype = sec.get("type", "table")

        if stype == "info_table":
            data = [[Paragraph(str(k), cell_style), Paragraph(str(v), cell_style)] for k, v in sec["rows"]]
            tbl = Table(data, colWidths=[55*mm, AVAIL_WIDTH - 55*mm], repeatRows=0)
            tbl.setStyle(TableStyle([
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#e3f2fd")),
                ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#bbbbbb")),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("TOPPADDING", (0, 0), (-1, -1), 4),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ]))
            story.append(tbl)

        elif stype == "table":
            if sec["headers"] and sec["rows"]:
                n_cols = len(sec["headers"])
                # 动态列宽：根据列数分配，总宽不超过可用宽度
                col_w = AVAIL_WIDTH / n_cols
                # 如果列很多，限制最小列宽并允许小字体
                font_size = 8
                if n_cols >= 12:
                    font_size = 6
                elif n_cols >= 8:
                    font_size = 7
                cell_style2 = ParagraphStyle("CellCN2", fontName=CJK_FONT, fontSize=font_size,
                                             leading=font_size + 2, alignment=TA_LEFT)
                hdr_style = ParagraphStyle("HdrCN", fontName=CJK_FONT, fontSize=font_size + 1,
                                           leading=font_size + 3, alignment=TA_CENTER,
                                           textColor=colors.white)
                # 表头
                header_cells = [Paragraph(str(h), hdr_style) for h in sec["headers"]]
                # 数据行（用Paragraph包装以自动换行）
                data_rows = []
                for r in sec["rows"]:
                    row_cells = []
                    for i, v in enumerate(r):
                        text = _fmt_cell(v)
                        row_cells.append(Paragraph(text, cell_style2))
                    data_rows.append(row_cells)
                data = [header_cells] + data_rows
                col_widths = [col_w] * n_cols
                tbl = Table(data, colWidths=col_widths, repeatRows=1)
                tbl.setStyle(TableStyle([
                    ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#1565c0")),
                    ("GRID", (0, 0), (-1, -1), 0.4, colors.HexColor("#bbbbbb")),
                    ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                    ("TOPPADDING", (0, 0), (-1, -1), 2),
                    ("BOTTOMPADDING", (0, 0), (-1, -1), 2),
                    ("LEFTPADDING", (0, 0), (-1, -1), 2),
                    ("RIGHTPADDING", (0, 0), (-1, -1), 2),
                ]))
                story.append(tbl)

        elif stype == "charts":
            for ch in sec["charts"]:
                img_data = ch["image"]
                img = Image(io.BytesIO(img_data), width=AVAIL_WIDTH, height=AVAIL_WIDTH * 0.4)
                story.append(img)
                story.append(Spacer(1, 6))

        story.append(Spacer(1, 10))

    doc.build(story)

# ============================================================
# Word生成（python-docx）
# ============================================================
def _write_word(filepath, content):
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from PIL import Image as PILImage
    import io as _io

    doc = Document()

    # 标题
    title = doc.add_heading(content["title"], level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"报告生成时间: {content['generated_at']}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    for sec in content["sections"]:
        doc.add_heading(sec["title"], level=1)
        stype = sec.get("type", "table")

        if stype == "info_table":
            table = doc.add_table(rows=0, cols=2)
            table.style = "Light Shading Accent 1"
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            for k, v in sec["rows"]:
                row = table.add_row()
                row.cells[0].text = str(k)
                row.cells[1].text = str(v)
                for cell in row.cells:
                    for par in cell.paragraphs:
                        for r in par.runs:
                            r.font.size = Pt(10)

        elif stype == "table":
            if sec["headers"] and sec["rows"]:
                table = doc.add_table(rows=1, cols=len(sec["headers"]))
                table.style = "Light Shading Accent 1"
                hdr = table.rows[0]
                for i, h in enumerate(sec["headers"]):
                    hdr.cells[i].text = str(h)
                for row_vals in sec["rows"]:
                    row = table.add_row()
                    for i, v in enumerate(row_vals):
                        if i < len(row.cells):
                            row.cells[i].text = str(v)
                for row in table.rows:
                    for cell in row.cells:
                        for par in cell.paragraphs:
                            for r in par.runs:
                                r.font.size = Pt(9)

        elif stype == "charts":
            for ch in sec["charts"]:
                img_bytes = ch["image"]
                try:
                    pil = PILImage.open(_io.BytesIO(img_bytes))
                    # 估算宽度以适应页面
                    w_cm = min(16, pil.width / 96 * 2.54)
                    h_cm = w_cm * pil.height / pil.width
                    stream = _io.BytesIO(img_bytes)
                    doc.add_picture(stream, width=Cm(w_cm))
                except Exception:
                    pass
                cap = doc.add_paragraph()
                cap_run = cap.add_run(ch["title"])
                cap_run.font.size = Pt(9)
                cap_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

        doc.add_paragraph()

    doc.save(filepath)


# ============================================================
# 功能3: 多工况批处理报告（IEC DLC载荷包络）
# ============================================================
def generate_batch_report(batch_result, fmt="pdf", output_dir=None,
                          calc_time=None, filename=None):
    """
    生成多工况批处理报告（含载荷包络结论置顶）
    batch_result: serialize_batch_result() 的返回值
    fmt: pdf / word
    output_dir: 保存目录（可选）
    calc_time: 计算耗时
    filename: 自定义文件名（含扩展名，可选）
    返回: (文件路径, 文件名)
    """
    output_dir = output_dir or os.path.join(os.path.dirname(os.path.abspath(__file__)), "output")
    os.makedirs(output_dir, exist_ok=True)
    ts = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = "垂直轴风轮载荷分析结果（多工况批处理）"

    content = _build_batch_content(batch_result, calc_time)

    if fmt.lower() in ("pdf",):
        ext = "pdf"
        if filename:
            fname = filename if filename.lower().endswith(".pdf") else filename + ".pdf"
        else:
            fname = f"{base_name}_{ts}.pdf"
        filepath = os.path.join(output_dir, fname)
        _write_pdf(filepath, content)
    elif fmt.lower() in ("word", "docx", "doc"):
        ext = "docx"
        if filename:
            fname = filename if filename.lower().endswith(".docx") else filename + ".docx"
        else:
            fname = f"{base_name}_{ts}.docx"
        filepath = os.path.join(output_dir, fname)
        _write_word(filepath, content)
    else:
        raise ValueError(f"不支持的格式: {fmt}")

    return filepath, fname


def _build_batch_content(batch_result, calc_time=None):
    """构建多工况批处理报告内容模型"""
    content = {
        "title": "垂直轴风轮载荷分析结果（多工况批处理）",
        "generated_at": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
        "calc_time": calc_time or "--",
        "sections": [],
    }

    env = batch_result.get("envelope", {}) or {}
    cases = batch_result.get("cases", [])

    # 输入文件信息
    info_rows = [
        ("工况数量", str(batch_result.get("num_cases", 0))),
        ("叶片数", str(batch_result.get("num_blades", 0))),
        ("叶片疲劳指数 m", str(env.get("blade_m", "--"))),
        ("塔顶疲劳指数 m", str(env.get("tower_m", "--"))),
    ]
    if cases:
        names = ", ".join(c.get("filename", "") for c in cases[:5])
        if len(cases) > 5:
            names += f" 等{len(cases)}个"
        info_rows.append(("工况文件", names))
    content["sections"].append({
        "title": "输入文件信息",
        "type": "info_table",
        "rows": info_rows,
    })

    # ===== 一、结论：载荷包络总览（置顶）=====
    conclusion_rows = []
    tower_env = env.get("tower", {})
    for r in (tower_env.get("extreme") or []):
        conclusion_rows.append((
            r.get("通道", ""),
            f"上限={r.get('上限', '--')} ({r.get('上限来源', '--')}@{r.get('上限时刻', '--')}s)",
            f"下限={r.get('下限', '--')} ({r.get('下限来源', '--')}@{r.get('下限时刻', '--')}s)",
            "--", "--", "--",
        ))
    for r in (tower_env.get("fatigue") or []):
        conclusion_rows.append((
            r.get("通道", "") + " 等效疲劳",
            f"{r.get('等效疲劳载荷', '--')} (m={r.get('Wöhler指数m', '--')})",
            "--", "--", "--", "--",
        ))
    if conclusion_rows:
        content["sections"].insert(0, {
            "title": "一、结论：塔顶载荷包络（极限+疲劳）",
            "type": "table",
            "headers": ["校核项", "上限(来源@时刻)", "下限(来源@时刻)", "", "", ""],
            "rows": conclusion_rows,
        })

    # ===== 叶片校核包络（逐叶片）=====
    blade_env = env.get("blade", {})
    for bid, benv in blade_env.items():
        rows = []
        for r in (benv.get("extreme") or []):
            rows.append((r.get("通道", ""),
                         f"上限={r.get('上限', '--')} ({r.get('上限来源', '--')}@{r.get('上限时刻', '--')}s)",
                         f"下限={r.get('下限', '--')} ({r.get('下限来源', '--')}@{r.get('下限时刻', '--')}s)"))
        for r in (benv.get("fatigue") or []):
            rows.append((r.get("通道", "") + " 等效疲劳(m=%s)" % r.get("Wöhler指数m", "--"),
                         f"{r.get('等效疲劳载荷', '--')}", "--"))
        if rows:
            content["sections"].append({
                "title": f"二、叶片校核载荷包络 - {bid}",
                "type": "table",
                "headers": ["校核项", "极限/疲劳值", "备注"],
                "rows": rows,
            })

    # ===== 各工况明细 =====
    for c in cases:
        fi_rows = [
            ("工况ID", str(c.get("case_id", "--"))),
            ("文件名", str(c.get("filename", "--"))),
            ("文件大小", f"{c.get('size_mb', 0):.2f} MB"),
            ("工况类型", str(c.get("condition_type", "--"))),
            ("权重", str(c.get("weight", "--"))),
            ("数据", f"{c.get('num_rows', 0)} 行 x {c.get('num_cols', 0)} 列"),
            ("叶片数", str(c.get("num_blades", "--"))),
            ("采样率", f"{c.get('sample_rate', 0)} Hz"),
        ]
        content["sections"].append({
            "title": f"工况明细: {c.get('case_id', '--')}",
            "type": "info_table",
            "rows": fi_rows,
        })

    # 计算信息
    content["sections"].append({
        "title": "计算信息",
        "type": "info_table",
        "rows": [
            ("报告生成时间", content["generated_at"]),
            ("计算耗时", str(content["calc_time"])),
        ],
    })

    return content
